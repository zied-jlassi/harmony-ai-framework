"""
DOCX Parser - 100% Local RGPD Compliant
Uses python-docx - NO cloud APIs
"""

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import json
from typing import Optional
from pathlib import Path


class DOCXParser:
    """Parse DOCX 100% local - RGPD compliant"""

    def __init__(self):
        pass

    def parse(self, file_path: str) -> dict:
        """
        Extract text, structure and tables from DOCX.

        Args:
            file_path: Path to DOCX file

        Returns:
            dict with metadata, sections and tables
        """
        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            raise ValueError(f"Invalid DOCX file: {file_path}")

        result = {
            "source": str(Path(file_path).name),
            "parser": "python-docx",
            "rgpd_compliant": True,
            "metadata": {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            },
            "content": {
                "sections": [],
                "tables": []
            }
        }

        # Extract paragraphs by heading structure
        current_section = {"heading": None, "level": 0, "content": []}

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""

            if style_name.startswith('Heading'):
                # Save previous section
                if current_section["content"]:
                    result["content"]["sections"].append(current_section)

                # Start new section
                level = 1
                if style_name[-1].isdigit():
                    level = int(style_name[-1])

                current_section = {
                    "heading": para.text,
                    "level": level,
                    "content": []
                }
            else:
                if para.text.strip():
                    current_section["content"].append(para.text)

        # Don't forget last section
        if current_section["content"]:
            result["content"]["sections"].append(current_section)

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                "index": table_idx,
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
                "data": []
            }

            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data["data"].append(row_data)

            result["content"]["tables"].append(table_data)

        return result

    def parse_to_json(self, file_path: str, output_path: Optional[str] = None) -> str:
        """
        Parse DOCX and save as JSON.

        Args:
            file_path: Path to DOCX file
            output_path: Optional output path for JSON

        Returns:
            JSON string
        """
        result = self.parse(file_path)
        json_str = json.dumps(result, ensure_ascii=False, indent=2)

        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(json_str)

        return json_str


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python docx_parser.py <file.docx> [output.json]")
        sys.exit(1)

    parser = DOCXParser()
    output = sys.argv[2] if len(sys.argv) > 2 else None
    result = parser.parse_to_json(sys.argv[1], output)
    print(result)
